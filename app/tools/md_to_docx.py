import html.parser
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import markdown

# Helper functions for styling tables and borders
def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "D3D3D3", "space": "0"},
        bottom={"sz": 12, "color": "00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

class HtmlToDocxParser(html.parser.HTMLParser):
    def __init__(self, document):
        super().__init__()
        self.doc = document
        self.current_paragraph = None
        self.bold = 0
        self.italic = 0
        self.code = 0
        self.pre = 0
        self.blockquote = 0
        self.heading_level = 0
        self.link_href = None
        self.mark = 0
        self.span_classes = []
        
        self.list_type = []  # Stack of 'ul' or 'ol'
        self.list_count = [] # Stack of list item counters (integers)
        
        self.table = None
        self.table_rows = []
        self.current_row = None
        self.in_header = False
        
        # Buffer for accumulating text within elements
        self.text_buffer = ""

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            self.heading_level = int(tag[1])
            self.start_paragraph()
        elif tag == 'p':
            if not self.table_rows and self.current_row is None:
                self.start_paragraph()
        elif tag in ('strong', 'b'):
            self.bold += 1
        elif tag in ('em', 'i'):
            self.italic += 1
        elif tag == 'code':
            self.code += 1
        elif tag == 'pre':
            self.pre += 1
            self.start_paragraph()
        elif tag == 'blockquote':
            self.blockquote += 1
        elif tag == 'ul':
            self.list_type.append('ul')
        elif tag == 'ol':
            self.list_type.append('ol')
            self.list_count.append(1)
        elif tag == 'li':
            self.start_paragraph()
        elif tag == 'table':
            self.table_rows = []
        elif tag == 'tr':
            self.current_row = []
        elif tag in ('th', 'td'):
            self.in_header = (tag == 'th')
            self.text_buffer = ""
        elif tag == 'a':
            self.link_href = attrs_dict.get('href')
        elif tag == 'mark':
            self.mark += 1
        elif tag == 'span':
            span_class = attrs_dict.get('class', '')
            self.span_classes.append(span_class)

    def handle_endtag(self, tag):
        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            self.heading_level = 0
            self.current_paragraph = None
        elif tag == 'p':
            self.current_paragraph = None
        elif tag in ('strong', 'b'):
            self.bold = max(0, self.bold - 1)
        elif tag in ('em', 'i'):
            self.italic = max(0, self.italic - 1)
        elif tag == 'code':
            self.code = max(0, self.code - 1)
        elif tag == 'pre':
            self.pre = max(0, self.pre - 1)
            self.current_paragraph = None
        elif tag == 'blockquote':
            self.blockquote = max(0, self.blockquote - 1)
        elif tag == 'ul':
            if self.list_type:
                self.list_type.pop()
        elif tag == 'ol':
            if self.list_type:
                self.list_type.pop()
            if self.list_count:
                self.list_count.pop()
        elif tag == 'li':
            self.current_paragraph = None
            if self.list_type and self.list_type[-1] == 'ol' and self.list_count:
                self.list_count[-1] += 1
        elif tag in ('th', 'td'):
            if self.current_row is not None:
                self.current_row.append((self.text_buffer, self.in_header))
            self.text_buffer = ""
        elif tag == 'tr':
            if self.current_row is not None:
                self.table_rows.append(self.current_row)
            self.current_row = None
        elif tag == 'table':
            self.render_table()
            self.table_rows = []
        elif tag == 'a':
            self.link_href = None
        elif tag == 'mark':
            self.mark = max(0, self.mark - 1)
        elif tag == 'span':
            if self.span_classes:
                self.span_classes.pop()

    def handle_data(self, data):
        if not data:
            return
        
        # If inside a table cell, accumulate data in buffer
        if self.current_row is not None:
            self.text_buffer += data
            return

        # Otherwise add run to paragraph
        if self.current_paragraph is None:
            # Skip empty whitespace between block tags
            if not data.strip():
                return
            self.start_paragraph()
            
        run = self.current_paragraph.add_run(data)
        self.style_run(run)

    def start_paragraph(self):
        if self.heading_level > 0:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
        elif self.list_type:
            p = self.doc.add_paragraph()
            level = len(self.list_type) - 1
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(0)
            
            # Add prefix manually
            prefix = ""
            if self.list_type[-1] == 'ul':
                bullets = ["•", "◦", "▪", "▫"]
                prefix = bullets[level % len(bullets)] + " "
            elif self.list_type[-1] == 'ol':
                count = self.list_count[-1] if self.list_count else 1
                prefix = f"{count}. "
                
            run = p.add_run(prefix)
            run.bold = True
            run.font.name = 'Inter'
            run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6) # Brand blue
        elif self.pre > 0:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add border & shading to pre blocks
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="D1D5DB"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F3F4F6"/>')
            p._p.get_or_add_pPr().append(shd)
        else:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            
            if self.blockquote > 0:
                p.paragraph_format.left_indent = Inches(0.4)
                # Left border for blockquote
                pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="12" w:color="3B82F6"/></w:pBdr>')
                p._p.get_or_add_pPr().append(pBdr)
                # Soft blue shading
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EFF6FF"/>')
                p._p.get_or_add_pPr().append(shd)
                
        self.current_paragraph = p

    def style_run(self, run):
        run.font.name = 'Inter'
        
        if self.heading_level > 0:
            run.bold = True
            if self.heading_level == 1:
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            elif self.heading_level == 2:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            else:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
        else:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            
            if self.bold > 0:
                run.bold = True
            if self.italic > 0:
                run.italic = True
            if self.code > 0 or self.pre > 0:
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
                if self.pre > 0:
                    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
                    if self.span_classes:
                        classes = self.span_classes[-1].split()
                        for c in classes:
                            # Comments
                            if c in ('c', 'c1', 'cm', 'cp', 'cs', 'ch', 'cpf'):
                                run.font.color.rgb = RGBColor(0x6A, 0x73, 0x7D)
                                run.italic = True
                                break
                            # Keywords
                            elif c in ('k', 'kd', 'kn', 'kp', 'kr', 'kt', 'kc'):
                                run.font.color.rgb = RGBColor(0xD7, 0x3A, 0x49)
                                run.bold = True
                                break
                            # Strings
                            elif c in ('s', 's1', 's2', 'sb', 'sc', 'sd', 'se', 'sh', 'si', 'sx', 'sr', 'ss'):
                                run.font.color.rgb = RGBColor(0x03, 0x2F, 0x62)
                                break
                            # Numbers
                            elif c in ('m', 'mb', 'mf', 'mh', 'mi', 'mo', 'il'):
                                run.font.color.rgb = RGBColor(0x00, 0x5C, 0xC5)
                                break
                            # Functions
                            elif c in ('nf', 'fm'):
                                run.font.color.rgb = RGBColor(0x6F, 0x42, 0xC1)
                                break
                            # Builtins
                            elif c in ('nb', 'bp'):
                                run.font.color.rgb = RGBColor(0xE3, 0x62, 0x09)
                                break
                            # Tags
                            elif c == 'nt':
                                run.font.color.rgb = RGBColor(0x22, 0x86, 0x3A)
                                break
                            # Attributes
                            elif c == 'na':
                                run.font.color.rgb = RGBColor(0x6F, 0x42, 0xC1)
                                break
                            # Operators and punctuation
                            elif c in ('o', 'p'):
                                run.font.color.rgb = RGBColor(0xD7, 0x3A, 0x49)
                                break
            if self.blockquote > 0:
                run.italic = True
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            if self.link_href:
                run.underline = True
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        
        if self.mark > 0:
            from docx.enum.text import WD_COLOR_INDEX
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
    def render_table(self):
        if not self.table_rows:
            return
        
        num_cols = max(len(row) for row in self.table_rows)
        num_rows = len(self.table_rows)
        
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for r_idx, row_data in enumerate(self.table_rows):
            row = table.rows[r_idx]
            for c_idx, cell_data in enumerate(row_data):
                if c_idx >= num_cols:
                    break
                text, is_header = cell_data
                cell = row.cells[c_idx]
                
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(text.strip())
                run.font.name = 'Inter'
                run.font.size = Pt(9.5)
                
                if is_header:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    set_cell_background(cell, "3B82F6")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    if r_idx % 2 == 0:
                        set_cell_background(cell, "F9FAFB")
                    else:
                        set_cell_background(cell, "FFFFFF")
                
                set_cell_border(
                    cell,
                    top={"sz": 4, "val": "single", "color": "E5E7EB"},
                    bottom={"sz": 4, "val": "single", "color": "E5E7EB"},
                    left={"sz": 4, "val": "single", "color": "E5E7EB"},
                    right={"sz": 4, "val": "single", "color": "E5E7EB"}
                )
        
        self.doc.add_paragraph().paragraph_format.space_before = Pt(6)

def convert_markdown_to_docx(md_text: str, output_path: str) -> None:
    # Convert markdown to html with fenced_code, codehilite and extra
    # Support ==highlight== in code blocks and text using unicode placeholders
    import re
    md_marked = re.sub(r'==(?!\s)([^=]+?)(?<!\s)==', '\ue000\\g<1>\ue001', md_text)
    html_content = markdown.markdown(md_marked, extensions=["fenced_code", "codehilite", "extra"])
    
    # Replace unicode placeholders back to HTML <mark> tags
    html_content = re.sub(r'(?:<span class="[^\"]+">)?\ue000(?:</span>)?', '<mark>', html_content)
    html_content = re.sub(r'(?:<span class="[^\"]+">)?\ue001(?:</span>)?', '</mark>', html_content)
    
    doc = Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Setup parser
    parser = HtmlToDocxParser(doc)
    parser.feed(html_content)
    
    # Save document
    doc.save(output_path)
